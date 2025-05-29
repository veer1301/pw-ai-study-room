import os
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from striprtf.striprtf import rtf_to_text



def extract_text(file_path, lang='eng+hin'):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_ext = file_path.rsplit('.', 1)[-1].lower()

    if file_ext == 'pdf':
        return extract_text_from_pdf(file_path, lang)
    elif file_ext in {'jpg', 'jpeg', 'png'}:
        return extract_text_from_image(file_path, lang)
    elif file_ext == 'txt':
        return extract_text_from_txt(file_path)
    elif file_ext == 'docx':
        return extract_text_from_docx(file_path)
    elif file_ext == 'rtf':
        return extract_text_from_rtf(file_path)
    elif file_ext == 'doc':
        raise NotImplementedError("DOC format not supported. Convert to DOCX or use an external tool.")
    else:
        raise ValueError("Unsupported file type. Use PDF, JPG, JPEG, PNG, TXT, DOCX, or RTF.")

def extract_text_from_image(image_path, lang):
    image = Image.open(image_path)
    return pytesseract.image_to_string(image, lang=lang)

def extract_text_from_pdf(pdf_path, lang):
    images = convert_from_path(pdf_path)
    text = ''
    for i, image in enumerate(images):
        text += f"\n--- Page {i + 1} ---\n"
        text += pytesseract.image_to_string(image, lang=lang)
    return text.strip()

def extract_text_from_txt(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(filter(None, full_text))

def extract_text_from_rtf(rtf_path):
    with open(rtf_path, 'r', encoding='utf-8') as f:
        return rtf_to_text(f.read())

if __name__ == "__main__":
    file_path = "requirements.txt"
    extracted_text = extract_text(file_path)
    print("Extracted Text:\n")
    print(extracted_text)